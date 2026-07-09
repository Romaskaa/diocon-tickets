from __future__ import annotations

from typing import TYPE_CHECKING

from html import escape
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

if TYPE_CHECKING:
    from .services.stage_export import ProjectStagesReport

HEADERS = [
    "№",
    "Этап",
    "Статус",
    "Плановое начало",
    "Плановое завершение",
    "Фактическое начало",
    "Фактическое завершение",
    "Ответственный",
    "Просрочен",
    "Длительность (дни)",
    "Описание",
    "Критерии завершения",
]

PDF_FONT_NAME = "ProjectReportFont"
PDF_BOLD_FONT_NAME = "ProjectReportFontBold"
PDF_TABLE_COLUMN_WIDTHS = [22, 70, 45, 58, 64, 64, 70, 80, 42, 58, 110, 110]


def export_project_stages_to_excel(report: ProjectStagesReport) -> bytes:
    """
    Сформировать Excel-файл отчёта по этапам проекта.
    """

    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Этапы проекта"

    worksheet["A1"] = f"Отчёт по этапам проекта: {report.project_name}"
    worksheet["A2"] = f"Ссылка на проект: {report.project_url}"
    worksheet["A2"].hyperlink = report.project_url
    worksheet["A2"].style = "Hyperlink"
    worksheet["A3"] = f"Статус проекта: {report.project_status}"
    worksheet["A4"] = f"Дата формирования: {report.generated_at:%d.%m.%Y %H:%M}"

    for cell in ("A1", "A2", "A3", "A4"):
        worksheet[cell].font = Font(bold=True)

    header_row = 6
    for column_index, header in enumerate(HEADERS, start=1):
        cell = worksheet.cell(row=header_row, column=column_index, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="4F81BD")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for row_index, row in enumerate(report.rows, start=header_row + 1):
        values = [
            row.number,
            row.name,
            row.status,
            row.planned_start,
            row.planned_end,
            row.started_at,
            row.completed_at,
            row.responsible,
            row.is_overdue,
            row.planned_duration_days,
            row.description,
            row.completion_criteria,
        ]

        for column_index, value in enumerate(values, start=1):
            cell = worksheet.cell(row=row_index, column=column_index, value=value)
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    worksheet.freeze_panes = "A7"
    worksheet.auto_filter.ref = (
        f"A{header_row}:"
        f"{get_column_letter(len(HEADERS))}{max(header_row, header_row + len(report.rows))}"
    )

    column_widths = [6, 28, 16, 18, 20, 20, 22, 38, 14, 18, 42, 42]
    for column_index, width in enumerate(column_widths, start=1):
        worksheet.column_dimensions[get_column_letter(column_index)].width = width

    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def export_project_stages_to_pdf(report: ProjectStagesReport) -> bytes:
    """
    Сформировать PDF-файл отчёта по этапам проекта.
    """

    regular_font, bold_font = _register_pdf_fonts()

    output = BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=landscape(A4),
        leftMargin=24,
        rightMargin=24,
        topMargin=24,
        bottomMargin=24,
    )

    styles = getSampleStyleSheet()
    styles["Title"].fontName = bold_font
    styles["Normal"].fontName = regular_font

    table_header_style = ParagraphStyle(
        name="ProjectStagesTableHeader",
        parent=styles["Normal"],
        fontName=bold_font,
        fontSize=6,
        leading=7,
        textColor=colors.white,
        alignment=1,
    )
    table_cell_style = ParagraphStyle(
        name="ProjectStagesTableCell",
        parent=styles["Normal"],
        fontName=regular_font,
        fontSize=6,
        leading=7,
        wordWrap="CJK",
    )

    project_url = escape(report.project_url)
    elements = [
        Paragraph(f"Отчёт по этапам проекта: {escape(report.project_name)}", styles["Title"]),
        Paragraph(
            f'Ссылка на проект: <a href="{project_url}">{project_url}</a>',
            styles["Normal"],
        ),
        Paragraph(f"Статус проекта: {escape(report.project_status)}", styles["Normal"]),
        Paragraph(f"Дата формирования: {report.generated_at:%d.%m.%Y %H:%M}", styles["Normal"]),
        Spacer(1, 12),
    ]

    data = [[_pdf_cell(header, table_header_style) for header in HEADERS]]
    data.extend(
        [
            _pdf_cell(row.number, table_cell_style),
            _pdf_cell(row.name, table_cell_style),
            _pdf_cell(row.status, table_cell_style),
            _pdf_cell(row.planned_start, table_cell_style),
            _pdf_cell(row.planned_end, table_cell_style),
            _pdf_cell(row.started_at, table_cell_style),
            _pdf_cell(row.completed_at, table_cell_style),
            _pdf_cell(row.responsible, table_cell_style),
            _pdf_cell(row.is_overdue, table_cell_style),
            _pdf_cell(row.planned_duration_days, table_cell_style),
            _pdf_cell(row.description, table_cell_style),
            _pdf_cell(row.completion_criteria, table_cell_style),
        ]
        for row in report.rows
    )

    table = Table(
        data,
        colWidths=PDF_TABLE_COLUMN_WIDTHS,
        hAlign="LEFT",
        repeatRows=1,
    )
    table.setStyle(
        TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4F81BD")),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 3),
            ("RIGHTPADDING", (0, 0), (-1, -1), 3),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ])
    )

    elements.append(table)
    document.build(elements)

    return output.getvalue()


def export_project_stages_to_word(report: ProjectStagesReport) -> bytes:
    """
    Сформировать Word-файл отчёта по этапам проекта.
    """

    document = Document()

    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)

    title = document.add_heading(f"Отчёт по этапам проекта: {report.project_name}", level=1)
    title.alignment = 1

    project_link_paragraph = document.add_paragraph("Ссылка на проект: ")
    _add_hyperlink(project_link_paragraph, report.project_url, report.project_url)

    document.add_paragraph(f"Статус проекта: {report.project_status}")
    document.add_paragraph(f"Дата формирования: {report.generated_at:%d.%m.%Y %H:%M}")

    table = document.add_table(rows=1, cols=len(HEADERS))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for index, header in enumerate(HEADERS):
        paragraph = header_cells[index].paragraphs[0]
        run = paragraph.add_run(header)
        run.bold = True
        run.font.size = Pt(8)

    for row in report.rows:
        cells = table.add_row().cells
        values = [
            row.number,
            row.name,
            row.status,
            row.planned_start,
            row.planned_end,
            row.started_at,
            row.completed_at,
            row.responsible,
            row.is_overdue,
            row.planned_duration_days,
            row.description,
            row.completion_criteria,
        ]

        for index, value in enumerate(values):
            paragraph = cells[index].paragraphs[0]
            run = paragraph.add_run(str(value))
            run.font.size = Pt(8)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _register_pdf_fonts() -> tuple[str, str]:
    regular_font_path = _find_font_file([
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibri.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf"),
    ])
    bold_font_path = _find_font_file([
        Path("C:/Windows/Fonts/arialbd.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
        Path("/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf"),
    ])

    if regular_font_path is None:
        return "Helvetica", "Helvetica-Bold"

    if PDF_FONT_NAME not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont(PDF_FONT_NAME, str(regular_font_path)))

    if bold_font_path is None:
        return PDF_FONT_NAME, PDF_FONT_NAME

    if PDF_BOLD_FONT_NAME not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont(PDF_BOLD_FONT_NAME, str(bold_font_path)))

    return PDF_FONT_NAME, PDF_BOLD_FONT_NAME


def _find_font_file(candidates: list[Path]) -> Path | None:
    return next((path for path in candidates if path.exists()), None)


def _pdf_cell(value: object, style: ParagraphStyle) -> Paragraph:
    text = escape(str(value)).replace("\n", "<br/>")
    return Paragraph(text, style)


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    new_run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_properties.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)

    new_run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)  # noqa: SLF001
